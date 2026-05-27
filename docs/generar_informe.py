from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors ────────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x0D, 0x4F, 0x8C)
BLUE_LIGHT = RGBColor(0x1A, 0x6B, 0xB5)
TEAL       = RGBColor(0x08, 0x91, 0xB2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK       = RGBColor(0x1E, 0x29, 0x3B)
GRAY       = RGBColor(0x47, 0x55, 0x69)
LIGHT_BG   = RGBColor(0xF1, 0xF5, 0xF9)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Default paragraph style ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_row_height(row, height_cm):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(height_cm * 567)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)

def add_paragraph(text='', bold=False, size=11, color=DARK, space_before=0,
                  space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.alignment    = align
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color
        run.font.name  = 'Calibri'
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(15)
    run.font.color.rgb = BLUE
    run.font.name  = 'Calibri'
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '0891B2')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(12)
    run.font.color.rgb = BLUE_LIGHT
    run.font.name  = 'Calibri'
    return p

def add_body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GRAY
    run.font.name  = 'Calibri'
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GRAY
    run.font.name  = 'Calibri'
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = GRAY
    run.font.name  = 'Calibri'
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    # shading
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F1F5F9')
    pPr.append(shd)
    return p

def add_info_box(text):
    """Light-teal info box"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'E0F2FE')
    cell.width = Inches(5.5)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('left',):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '18')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '0891B2')
        tcBdr.append(b)
    tcPr.append(tcBdr)
    p    = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run  = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = DARK
    run.font.name  = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_data_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, '0D4F8C')
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size      = Pt(10)
        run.font.name      = 'Calibri'
    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = 'F8FAFC' if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            run  = p.add_run(val)
            run.font.size      = Pt(10)
            run.font.color.rgb = GRAY
            run.font.name      = 'Calibri'
    # Col widths
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
# Blue cover block via 1-cell table
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cover_tbl.cell(0, 0)
set_cell_bg(cell, '0D4F8C')
set_row_height(cover_tbl.rows[0], 8)

def cover_line(cell, text, size, bold=False, italic=False, space_before=0):
    p   = cell.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size       = Pt(size)
    run.font.color.rgb  = WHITE
    run.font.name       = 'Calibri'
    return p

# Remove default empty paragraph in cell
cell.paragraphs[0].clear()
cover_line(cell, '', 6, space_before=20)
cover_line(cell, '🌊  Azul Horizonte', 28, bold=True)
cover_line(cell, 'Boutique Hotel', 14)
cover_line(cell, '"Donde el Pacífico te abraza"', 11, italic=True)
cover_line(cell, '', 6)
cover_line(cell, 'INFORME TÉCNICO DEL PROYECTO', 16, bold=True)
cover_line(cell, '', 6)
cover_line(cell, 'Universidad Mariano Gálvez de Guatemala', 10)
cover_line(cell, 'Inteligencia Artificial — Proyecto Final', 10)
cover_line(cell, 'Mayo 2026', 10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('1. Resumen Ejecutivo')

add_info_box(
    'Se construyó un sitio web completo para el Azul Horizonte Boutique Hotel '
    '(Playa Champerico, Retalhuleu) que integra un agente de IA conversacional '
    '(Claude) con un sistema automatizado de reservaciones orquestado por Zapier y n8n.'
)

add_body('El proyecto cubre tres dimensiones principales:')
add_numbered('Sitio web público — páginas estáticas de presentación del hotel, habitaciones, servicios y preguntas frecuentes.')
add_numbered('Agente de IA — chat widget embebido que responde preguntas con contexto del hotel (RAG) y captura reservaciones en tiempo real verificando disponibilidad en Google Calendar.')
add_numbered('Automatización post-reserva — flujos en Zapier y n8n que envían correos de confirmación, recordatorios 24 h y 2 h antes del check-in, gestionan cancelaciones/modificaciones y generan reportes mensuales al administrador.')

# ══════════════════════════════════════════════════════════════════════════════
# 2. STACK TECNOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('2. Stack Tecnológico')

add_data_table(
    ['Capa', 'Tecnología', 'Rol'],
    [
        ['Frontend / Backend',  'TanStack Start (React SSR)',         'Páginas web + API routes server-side'],
        ['Lenguaje',            'TypeScript',                          'Todo el codebase'],
        ['Estilos',             'Tailwind CSS',                        'UI del sitio y del chat widget'],
        ['Contenedores',        'Docker + docker-compose',             'Empaquetado y despliegue'],
        ['Despliegue',          'Azure Container Apps',                'Producción en la nube'],
        ['Agente IA',           'Claude API (claude-sonnet-4-6)',      'Chat conversacional + tool use'],
        ['Base de conocimiento','RAG por keywords sobre archivos .md', 'Contexto del hotel inyectado al agente'],
        ['Disponibilidad',      'Google Calendar API (Service Acct.)', 'Verifica si la suite está libre'],
        ['Automatización 1',    'Zapier',                              'Crear evento, enviar correo, registrar en Sheets'],
        ['Automatización 2',    'n8n (contenedor Docker)',             'Recordatorios, cancelaciones, reportes'],
        ['Base de datos',       'Google Sheets',                       'Registro de todas las reservaciones'],
        ['Correo',              'Gmail (vía Zapier y n8n)',            'Notificaciones al huésped y admin'],
    ],
    col_widths=[4, 5.5, 6.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. COMPONENTES PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('3. Componentes Principales del Sistema')

add_heading3('3.1 Sitio Web Público')
add_body('Cuatro páginas construidas con TanStack Start (React SSR):')
add_bullet('Home — propuesta de valor y llamadas a la acción.')
add_bullet('Habitaciones — tipos, precios y capacidad de las 3 suites.')
add_bullet('Servicios — experiencias y actividades del hotel.')
add_bullet('FAQ — preguntas frecuentes.')
add_body('El chat widget (botón flotante) está montado en el layout global y es visible en todas las páginas.')

add_heading3('3.2 Habitaciones Disponibles')
add_data_table(
    ['Código', 'Tipo', 'Precio / noche', 'Capacidad'],
    [
        ['SVM', 'Suite Vista Mar',       'Q 950  (~$120)', '1–2 personas'],
        ['SFM', 'Suite Frente al Mar',   'Q 1,450 (~$180)', '2–3 personas'],
        ['SP',  'Suite Presidencial',    'Q 2,200 (~$280)', '2–6 personas'],
    ],
    col_widths=[2.5, 5, 4, 4]
)

add_heading3('3.3 Chat Widget y Agente de IA')
add_body('El chat es un panel lateral deslizante. Cada mensaje del usuario pasa por dos pasos antes de llegar a Claude:')
add_numbered('RAG: rag.ts busca por keywords en los archivos .md del hotel y adjunta los fragmentos relevantes al system prompt.')
add_numbered('Loop agéntico en chatFn.ts: envía el historial completo + contexto RAG y ejecuta hasta 5 iteraciones de tool use para verificar disponibilidad y/o mostrar el formulario.')
add_body('El frontend muestra los mensajes intermedios ("Déjame verificar...") y la respuesta final como burbujas separadas, reflejando cada turno del loop agéntico.')

add_heading3('3.4 Módulo de Disponibilidad (google-calendar.ts)')
add_body(
    'La función isRoomAvailable(roomType, checkIn, checkOut) se comunica con Google Calendar API '
    'usando una Service Account. Busca eventos en el calendario único "hotel" cuyo título contenga '
    'el código de la suite y verifica que no haya traslape de fechas.'
)

add_heading3('3.5 Procesamiento de Reservaciones (bookFn.ts)')
add_numbered('Recibe el payload del formulario: nombre, email, teléfono, fechas, tipo de habitación, huéspedes.')
add_numbered('Genera un bookingId (UUID).')
add_numbered('Llama isRoomAvailable() → Google Calendar API.')
add_numbered('Si disponible: llama sendToZapier() y retorna { status: "confirmed", bookingId }.')
add_numbered('Si no disponible: retorna { status: "unavailable" } sin llamar Zapier.')

# ══════════════════════════════════════════════════════════════════════════════
# 4. FLUJO COMPLETO
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('4. Flujo Completo de una Reservación')

add_code_block(
"""Browser (Chat Widget)
  │
  ├─ Usuario: "Quiero reservar suite frente al mar del 27 al 28 de mayo"
  │
  ▼
POST /api/chat ──► chatFn.ts (Loop agéntico, máx. 5 iteraciones)
  │
  │  TURNO 1: Claude genera texto + tool_use check_availability
  │            └── "Déjame verificar la disponibilidad..."
  │                    │
  │            chatFn llama isRoomAvailable() → Google Calendar API
  │                    │
  │  TURNO 2: Claude recibe resultado y genera respuesta final
  │
  ▼
ChatResponse {
  checkingText: "Déjame verificar...",   ← mensaje 1 en el chat
  text:         "¡Excelente noticia!..."  ← mensaje 2 en el chat
}
  │
  ▼
Frontend añade ambos mensajes como burbujas separadas
  │
  │  Usuario confirma → Claude llama tool_use show_booking_form
  │
  ▼
Frontend renderiza <BookingForm /> dentro del chat
  │
  │  Usuario llena: nombre, email, teléfono, fechas, suite, huéspedes
  │
  ▼
bookServerFn (server-side)
  │
  ├─ isRoomAvailable() → Google Calendar API (segunda verificación al confirmar)
  │         │
  │    ┌────┴────┐
  │   Libre   Ocupada
  │    │          └── { status: "unavailable" } → mensaje en chat
  │    ▼
  sendToZapier({ bookingId, name, email, ... })
  │
  ▼
{ status: "confirmed", bookingId } → Frontend: pantalla de confirmación"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. ZAPIER
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('5. Configuración de Zapier')

add_info_box(
    'Trigger: Webhooks by Zapier → Catch Hook. El backend llama este webhook únicamente '
    'cuando la disponibilidad ya fue confirmada por Google Calendar. Zapier nunca recibe '
    'una reservación no disponible.'
)

add_heading3('Pasos del Zap (flujo lineal)')

add_numbered('Trigger — Catch Hook\nRecibe JSON: bookingId, name, email, phone, checkIn, checkOut, roomType (SVM/SFM/SP), guests. Fechas en ISO 8601 con offset Guatemala (2026-06-10T15:00:00-06:00).')
add_numbered('Formatter — Lookup Table\nSVM → Suite Vista Mar | SFM → Suite Frente al Mar | SP → Suite Presidencial')
add_numbered('Formatter — Date/Time: formatea checkIn → DD/MM/YYYY')
add_numbered('Formatter — Date/Time: formatea checkOut → DD/MM/YYYY')
add_numbered('Google Calendar → Create Detailed Event en el calendario "hotel".\nTítulo: Reservación - {roomType} - {name} | Zona horaria: America/Guatemala')
add_numbered('Gmail → Send Email al huésped con correo de confirmación (usa nombre completo de la suite y fechas formateadas).')
add_numbered('Google Sheets → Create Spreadsheet Row en la hoja "Reservaciones" — último paso del flujo.')

add_heading3('Esquema de la hoja "Reservaciones" (Google Sheets)')
add_data_table(
    ['Columna', 'Tipo', 'Notas'],
    [
        ['booking_id',        'string',   'UUID generado por el backend'],
        ['name',              'string',   ''],
        ['email',             'string',   ''],
        ['phone',             'string',   ''],
        ['check_in',          'date',     'YYYY-MM-DD'],
        ['check_out',         'date',     'YYYY-MM-DD'],
        ['room_type',         'string',   'SVM / SFM / SP'],
        ['guests',            'number',   ''],
        ['status',            'string',   'confirmed / cancelled / modified'],
        ['calendar_event_id', 'string',   'ID del evento en Google Calendar (guardado por Zapier)'],
        ['reminder_24h',      'boolean',  'Flag para recordatorio 24 h'],
        ['reminder_2h',       'boolean',  'Flag para recordatorio 2 h'],
        ['created_at',        'datetime', ''],
    ],
    col_widths=[4, 3, 9]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. n8n WORKFLOWS
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('6. Workflows de n8n')

add_info_box(
    'n8n corre como segundo contenedor en docker-compose (puerto 5678). Los workflows se '
    'importan desde los archivos JSON en n8n-workflows/. Cubren el ciclo de vida completo '
    'de la reservación: recordatorios, cancelación, modificación y reportes al administrador.'
)

add_heading3('Workflow A — Recordatorio 24 horas antes  (reminder-24h.json)')
add_data_table(
    ['Elemento', 'Detalle'],
    [
        ['Trigger',  'Schedule cron 0 14 * * * (8:00 AM Guatemala, UTC-6)'],
        ['Paso 1',   'Google Sheets → leer todas las filas de "Reservaciones"'],
        ['Paso 2',   'Code → filtrar: check_in == mañana AND status == confirmed AND reminder_24h != TRUE'],
        ['Paso 3',   'Gmail → enviar recordatorio con link de cancelación (N8N_HOST/webhook/cancel-booking?id=...)'],
        ['Paso 4',   'Google Sheets → actualizar reminder_24h = TRUE'],
    ],
    col_widths=[3, 13]
)

add_heading3('Workflow B — Recordatorio 2 horas antes  (reminder-2h.json)')
add_data_table(
    ['Elemento', 'Detalle'],
    [
        ['Trigger',  'Schedule: cada 15 minutos'],
        ['Paso 1',   'Google Sheets → leer todas las filas'],
        ['Paso 2',   'Code → filtrar: status == confirmed AND reminder_2h != TRUE AND check_in ∈ [ahora, ahora+2h]'],
        ['Paso 3',   'Gmail → enviar "Tu check-in es en 2 horas"'],
        ['Paso 4',   'Google Sheets → actualizar reminder_2h = TRUE'],
        ['Nota',     'El flag reminder_2h garantiza exactamente un correo por reservación sin importar cuántas veces se ejecute el trigger en la ventana.'],
    ],
    col_widths=[3, 13]
)

add_heading3('Workflow C — Cancelación de reserva  (cancelacion.json)')
add_data_table(
    ['Elemento', 'Detalle'],
    [
        ['Trigger',     'Webhook GET /webhook/cancel-booking?id={booking_id} (link en el email de recordatorio 24h)'],
        ['Paso 1',      'Google Sheets → buscar fila por booking_id'],
        ['Paso 2',      'IF found == true AND status == confirmed'],
        ['Rama TRUE',   'Google Calendar → eliminar evento → Sheets (status=cancelled) → Gmail → HTML "Cancelación exitosa"'],
        ['Rama FALSE',  'HTML: "Enlace no válido o ya procesado"'],
    ],
    col_widths=[3, 13]
)

add_heading3('Workflow D — Modificación de fechas  (modificacion.json)')
add_data_table(
    ['Elemento', 'Detalle'],
    [
        ['Trigger',   'Form Trigger en /form/modify-booking (campos: ID de Reservación, nueva fecha check-in, nueva fecha check-out)'],
        ['Paso 1',    'Code → validar formato de fechas ISO 8601 con offset Guatemala'],
        ['Paso 2',    'Google Sheets → buscar la reservación por booking_id'],
        ['Paso 3',    'IF found == true AND status == confirmed'],
        ['Rama TRUE', 'Calendar (eliminar evento anterior) → Calendar (crear nuevo evento) → Sheets (actualizar fechas, calendar_event_id, status=modified, resetear flags) → Gmail'],
        ['Rama FALSE','Termina sin acción'],
    ],
    col_widths=[3, 13]
)

add_heading3('Workflow E — Reporte mensual  (reporte-mensual.json)')
add_data_table(
    ['Elemento', 'Detalle'],
    [
        ['Trigger', 'Schedule cron 0 13 1 * * (7:00 AM Guatemala, día 1 de cada mes)'],
        ['Paso 1',  'Google Sheets → leer todas las filas'],
        ['Paso 2',  'Code → filtrar reservas del mes anterior; calcular totales por suite, noches e ingresos (Q)'],
        ['Paso 3',  'Gmail → enviar tabla HTML con estadísticas a ADMIN_EMAIL'],
    ],
    col_widths=[3, 13]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. AGENTE DE IA
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('7. Configuración del Agente de IA')

add_heading3('7.1 Modelo')
add_body('Se usa claude-haiku-4-5-20251001 (Anthropic Claude API). El agente corre enteramente en el servidor; la API key nunca llega al navegador del usuario.')

add_heading3('7.2 RAG (Retrieval-Augmented Generation)')
add_body('El módulo rag.ts implementa un RAG liviano sin base de datos vectorial:')
add_numbered('Los archivos fuente son 4 documentos Markdown en src/lib/knowledge/: rooms.md, policies.md, services.md, faq.md.')
add_numbered('Cada mensaje del usuario dispara retrieveContext(query), que hace matching por keywords sobre el texto de los .md.')
add_numbered('Los fragmentos más relevantes se adjuntan al system prompt antes de cada llamada a Claude.')
add_numbered('Para actualizar el conocimiento del agente basta con editar los archivos .md — no hay índice ni embedding que reconstruir.')

add_heading3('7.3 Tools del Agente')
add_body('El agente tiene dos tools definidos en claude.ts:')
add_data_table(
    ['Tool', 'Cuándo se invoca', 'Resultado'],
    [
        ['check_availability',
         'Cuando el cliente menciona fechas y tipo de suite, ANTES de mostrar el formulario.',
         'El backend llama isRoomAvailable() → Google Calendar API y devuelve disponible / no disponible.'],
        ['show_booking_form',
         'Solo después de confirmar disponibilidad con check_availability y cuando el cliente acepta proceder.',
         'El frontend renderiza <BookingForm /> dentro del historial del chat.'],
    ],
    col_widths=[4, 6, 6]
)

add_code_block(
"""// Tool 1 — verifica disponibilidad real en Google Calendar
{
  name: "check_availability",
  description: "Verifica si una habitación está disponible para las fechas solicitadas...",
  input_schema: {
    type: "object",
    properties: {
      roomType: { type: "string", enum: ["SVM", "SFM", "SP"] },
      checkIn:  { type: "string", description: "ISO 8601 offset Guatemala, ej: 2026-06-10T15:00:00-06:00" },
      checkOut: { type: "string", description: "ISO 8601 offset Guatemala, ej: 2026-06-12T12:00:00-06:00" }
    },
    required: ["roomType", "checkIn", "checkOut"]
  }
}

// Tool 2 — muestra el formulario HTML en el chat
{
  name: "show_booking_form",
  description: "Muestra el formulario de reservación. Usar SOLO después de verificar disponibilidad.",
  input_schema: {
    type: "object",
    properties: {
      roomType: { type: "string", enum: ["SVM", "SFM", "SP"] },
      checkIn:  { type: "string" },
      checkOut: { type: "string" }
    },
    required: ["roomType", "checkIn", "checkOut"]
  }
}"""
)

add_heading3('7.4 Loop Agéntico')
add_body('La función chat() en claude.ts implementa un loop de hasta 5 iteraciones:')
add_numbered('Llama a Claude API con el historial + context RAG.')
add_numbered('Si Claude invoca check_availability: el backend ejecuta isRoomAvailable(), devuelve el resultado como tool_result y vuelve al paso 1.')
add_numbered('Si Claude invoca show_booking_form: el backend marca showBookingForm=true y termina el loop.')
add_numbered('Si Claude solo responde texto (end_turn): termina el loop.')
add_body('El texto generado en el turno con check_availability se devuelve como checkingText (mensaje "Déjame verificar..."). El texto del turno final se devuelve como text. El frontend los muestra como dos burbujas separadas.')

add_heading3('7.5 Historial de conversación')
add_body('El historial se mantiene en estado React del cliente (useState) y se envía completo al backend en cada turno. No se persiste en base de datos.')

# ══════════════════════════════════════════════════════════════════════════════
# 8. POLÍTICAS DEL HOTEL
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('8. Políticas del Hotel')

add_data_table(
    ['Política', 'Detalle'],
    [
        ['Check-in / Check-out',    '3:00 PM entrada · 12:00 PM salida'],
        ['Estancia mínima',         '2 noches'],
        ['Cancelación',             'Gratuita con +48 h de anticipación; 50 % del total con menos de 48 h; 100 % de la primera noche por no-show'],
        ['Modificaciones',          'Sin costo con 72 h de anticipación, sujeto a disponibilidad'],
        ['Depósito',                '50 % del total al confirmar; saldo restante al check-in'],
        ['Pago',                    'Efectivo (Q o USD), Visa, Mastercard — no se aceptan cheques'],
        ['Mascotas / Eventos',      'No se permiten mascotas ni fiestas sin autorización'],
    ],
    col_widths=[5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. VARIABLES DE ENTORNO
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('9. Variables de Entorno')

add_heading3('App principal (.env)')
add_data_table(
    ['Variable', 'Descripción'],
    [
        ['ANTHROPIC_API_KEY',          'Clave de la Claude API'],
        ['ZAPIER_WEBHOOK_URL',          'URL del Catch Hook de Zapier'],
        ['GOOGLE_SERVICE_ACCOUNT_KEY', 'JSON de la Service Account (una sola línea)'],
        ['GOOGLE_CALENDAR_HOTEL',      'ID del calendario "hotel" en Google Calendar'],
        ['N8N_PASSWORD',               'Contraseña de acceso a la UI de n8n'],
    ],
    col_widths=[6, 10]
)

add_heading3('n8n (Settings → Environment Variables dentro de n8n)')
add_data_table(
    ['Variable', 'Descripción'],
    [
        ['GOOGLE_CALENDAR_ID', 'Mismo valor que GOOGLE_CALENDAR_HOTEL'],
        ['GOOGLE_SHEETS_ID',   'ID del spreadsheet (visible en la URL de Google Sheets)'],
        ['ADMIN_EMAIL',        'Correo del administrador para reportes mensuales'],
        ['N8N_HOST',           'URL pública de n8n (http://localhost:5678 local)'],
    ],
    col_widths=[5, 11]
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. DESPLIEGUE
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('10. Despliegue')

add_heading3('Desarrollo local')
add_code_block('# Sin Docker\nnpm run dev\n\n# Con Docker Compose (app + n8n)\ndocker-compose up --build')

add_heading3('Producción (Azure Container Apps)')
add_numbered('Construir imagen: docker build -t azul-horizonte-app .')
add_numbered('Push a Azure Container Registry: az acr build ...')
add_numbered('Crear Container App Environment en Azure Portal.')
add_numbered('Deploy del contenedor app (ingress público, puerto 3000).')
add_numbered('Deploy del contenedor n8n (ingress interno, puerto 5678).')
add_numbered('Configurar variables de entorno vía Azure Portal o az containerapp update --set-env-vars.')

# ══════════════════════════════════════════════════════════════════════════════
# 10. ESTRUCTURA DE ARCHIVOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading2('11. Estructura de Archivos Relevantes')

add_code_block(
"""/
├── src/
│   ├── routes/
│   │   ├── __root.tsx          # Layout global; monta <ChatWidget />
│   │   ├── index.tsx           # Home
│   │   ├── habitaciones.tsx    # Habitaciones y precios
│   │   ├── servicios.tsx       # Servicios
│   │   └── faq.tsx             # Preguntas frecuentes
│   ├── components/
│   │   ├── chat/
│   │   │   ├── ChatWidget.tsx  # Botón flotante + panel deslizante
│   │   │   ├── ChatMessage.tsx # Renderiza texto y BookingForm
│   │   │   └── BookingForm.tsx # Formulario embebido en el chat
│   │   └── layout/
│   │       ├── Header.tsx
│   │       └── Footer.tsx
│   └── lib/
│       ├── claude.ts           # buildSystemPrompt(), definición de tools
│       ├── rag.ts              # retrieveContext(query)
│       ├── zapier.ts           # sendToZapier()
│       ├── google-calendar.ts  # isRoomAvailable()
│       ├── server/
│       │   ├── bookFn.ts       # Verifica Calendar → llama Zapier
│       │   └── chatFn.ts       # Llama Claude API
│       └── knowledge/
│           ├── rooms.md
│           ├── policies.md
│           ├── services.md
│           └── faq.md
├── n8n-workflows/              # JSONs importables en n8n
├── Dockerfile
├── docker-compose.yml
└── .env.example"""
)

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
out = r'C:\Users\samue\OneDrive - Universidad Mariano Gálvez\Escritorio\IA proyecto final\docs\informe-proyecto.docx'
doc.save(out)
print(f'Saved: {out}')
